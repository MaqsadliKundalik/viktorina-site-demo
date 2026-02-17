from docx import Document
import sys

docx_path = "ATXM test.docx"
try:
    doc = Document(docx_path)
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    
    # Read first 50 paragraphs
    for i, para in enumerate(doc.paragraphs[:50]):
        text = para.text.strip()
        if text:
            print(f"--- Para {i+1} ---")
            print(text)
            
            # Check for bold text which might indicate correct answer
            for run in para.runs:
                if run.bold:
                    print(f"  [BOLD]: {run.text}")
            print("-----------------")
        
except Exception as e:
    print(f"Error reading DOCX: {e}")
