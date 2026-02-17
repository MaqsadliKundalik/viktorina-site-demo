from docx import Document
import sys

docx_path = "100 talik  test savollari ISTAT sirtqi 5-kurs.docx"

try:
    doc = Document(docx_path)
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    
    # Check first 20 paragraphs for formatting
    print("\n--- Examining Formatting (First 20 paragraphs) ---")
    for i, para in enumerate(doc.paragraphs[:20]):
        text = para.text.strip()
        if not text:
            continue
            
        print(f"\nPara {i+1}: {text[:50]}...")
        for run in para.runs:
            if run.bold:
                print(f"  [BOLD]: {run.text}")
            if run.font.color and run.font.color.rgb:
                 print(f"  [COLOR {run.font.color.rgb}]: {run.text}")

    # Check the last 10 paragraphs for answer keys
    print("\n--- Examining/Last 10 paragraphs ---")
    for i, para in enumerate(doc.paragraphs[-10:]):
        print(f"Para {len(doc.paragraphs)-10+i+1}: {para.text.strip()}")

except Exception as e:
    print(f"Error reading DOCX: {e}")
