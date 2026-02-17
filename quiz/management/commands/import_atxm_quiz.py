from django.core.management.base import BaseCommand
from quiz.models import Quiz, Question, Answer, Category
from docx import Document

class Command(BaseCommand):
    help = "ATXM DOCX fayldan testlarni import qiladi (++++ va ==== ajratgichlar bilan)"

    def add_arguments(self, parser):
        parser.add_argument("filepath", type=str, help="DOCX fayl yo'li")

    def handle(self, *args, **options):
        filepath = options["filepath"]
        
        try:
            doc = Document(filepath)
        except Exception as e:
            self.stdout.write(self.style.ERROR(f"DOCX o'qishda xatolik: {e}"))
            return

        base_category_name = "ATXM Testlari"
        base_quiz_title = "ATXM Quiz"
        category_counter = 1

        # Kategoriya yaratish
        current_category, _ = Category.objects.get_or_create(
            name=f"{base_category_name} {category_counter}"
        )
        current_quiz, _ = Quiz.objects.get_or_create(
            title=f"{base_quiz_title} {category_counter}",
            category=current_category
        )

        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text.strip())
        
        # Barcha matnni birlashtiramiz, keyin ajratamiz (bu oddiyroq bo'lishi mumkin, 
        # lekin paragraf chegaralari muhim bo'lsa, ehtiyot bo'lish kerak)
        
        # Yaxshiroq usul: Paragraflarni o'qib, ++++ va ==== ga qarab bo'laklash
        
        questions_blocks = []
        current_block = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text == "++++":
                if current_block:
                    questions_blocks.append(current_block)
                current_block = []
            else:
                if text: # Bo'sh qatorlarni tashlab ketamiz, lekin logikaga qarab
                    current_block.append(text)
        
        # Oxirgi blokni qo'shish
        if current_block:
            questions_blocks.append(current_block)

        self.stdout.write(f"Topilgan bloklar soni: {len(questions_blocks)}")

        imported = 0
        
        for block in questions_blocks:
            # Blok ichida ==== bilan ajratish
            # Blok bu stringlar ro'yxati
            
            # Stringlarni bitta matnga aylantirib, keyin ==== bo'yicha ajratamiz? 
            # Yo'q, chunki savol yoki javob ko'p qatorli bo'lishi mumkin.
            # Shuning uchun ==== separator sifatida ishlatiladi.
            
            parts = []
            current_part = []
            
            for line in block:
                if line == "====":
                    if current_part:
                        parts.append("\n".join(current_part).strip())
                    current_part = []
                else:
                    current_part.append(line)
            
            # Oxirgi qism
            if current_part:
                parts.append("\n".join(current_part).strip())
            
            if not parts:
                continue
                
            question_text = parts[0]
            answers_raw = parts[1:]
            
            if not question_text or not answers_raw:
                continue

            # Savollar soni 50 dan oshsa
            question_count = Question.objects.filter(quiz__category=current_category).count()
            if question_count >= 50:
                category_counter += 1
                current_category, _ = Category.objects.get_or_create(
                    name=f"{base_category_name} {category_counter}"
                )
                current_quiz, _ = Quiz.objects.get_or_create(
                    title=f"{base_quiz_title} {category_counter}",
                    category=current_category
                )
                self.stdout.write(self.style.WARNING(
                    f"⚠️ 50 savoldan oshdi, yangi kategoriya: {base_category_name} {category_counter}"
                ))

            question = Question.objects.create(
                quiz=current_quiz,
                text=question_text
            )
            
            for ans_text in answers_raw:
                is_correct = False
                if ans_text.startswith("#"):
                    is_correct = True
                    ans_text = ans_text[1:].strip()
                
                Answer.objects.create(
                    question=question,
                    text=ans_text,
                    is_correct=is_correct
                )
            imported += 1

        self.stdout.write(self.style.SUCCESS(f"✅ Import tugadi! {imported} ta savol qo'shildi."))
